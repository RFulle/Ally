# This code was made by Rhys Fuller for the AFJROTC UNIT IL-961
# The purpose of the code is to automate the production of promotion certificates.
# Version 1.0.0 was started at 4/24/2019
# Much of this code was repurposed from StackOverflow and other open-source websites. I'd also like to thank my friends Arick and Zachary for their help in giving me ideas on how to work around problems.
'''
MAKE PACKAGED VERSION
'''
import os
import comtypes.client
import time
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PyPDF2 import PdfFileMerger
cwd = os.getcwd()
def docxtopdf(path):
       wdFormatPDF = 17
       # absolute path is needed
       # be careful about the slash '\', use '\\' or '/' or raw string r"..."
       in_file = cwd + '/{}'.format(path)
       path2 = path.replace('.docx', '.pdf')
       out_file = cwd + '/{}'.format(path2)
       #in_file2 = r'absolute path of input docx file 2'
       #out_file2 = r'absolute path of outputpdf file 2'
       # print out filenames
       print("RENDERING DOCUMENTS:")
       print(in_file)
       print(out_file)
       #print(in_file2)
       #print(out_file2)
       # create COM object
       word = comtypes.client.CreateObject('Word.Application')
       print("25%")
       # key point 1: make word visible before open a new document
       word.Visible = True
       print("40%")
       # key point 2: wait for the COM Server to prepare well.
       time.sleep(3)
       print("COM SERVER WAIT COMPLETED")
       # convert docx file 1 to pdf file 1
       doc=word.Documents.Open(in_file) # open docx file 1
       print("60%")
       doc.SaveAs(out_file, FileFormat=wdFormatPDF) # conversion
       print("85%")
       doc.Close() # close docx file 1
       print("100%")
       print('Complete_')
       pdfs.append(out_file)
       word.Visible = False
       word.Quit() # close Word Application
def docmaker(name, rank):
       doc = Document()
       section = doc.sections[-1]
       section.orientation = WD_ORIENT.LANDSCAPE
       sections = doc.sections
       margin = 0.5
       for section in sections:
              section.top_margin = Inches(margin)
              section.bottom_margin = Inches(margin)
              section.left_margin = Inches(margin)
              section.right_margin = Inches(margin)
       style = doc.styles['Normal']
       font = style.font
       font.name = 'Times New Roman'
       font.size = Pt(14)
       new_width, new_height = section.page_height, section.page_width
       section.orientation = WD_ORIENT.LANDSCAPE
       section.page_width = new_width
       section.page_height = new_height
#code above sets page as landscape and sets margins, plus style "normal"
       doc.add_picture(cwd + '/AfJROTC EMBLEM.gif', width=Inches(1.75))
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#Code above sets picture
       cert = doc.add_paragraph()
       tim = cert.add_run('Certificate of Promotion')
       tim.bold = True
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
       font = tim.font
       font.size = Pt(24)
       font.color.rgb = RGBColor(31, 73, 125)
#That adds the "Certificate of promotion" part
       rec = doc.add_paragraph()
       bob = rec.add_run('IN RECOGNITION OF DEMONSTRATED FIDELITY AND ABILITIES I HEREBY APPOINT')
       bob.bold = True
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#Above is the "In recognition..." part
       sorry = doc.add_paragraph()
       geo = sorry.add_run(name)
       geo.bold = True
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
       font = geo.font
       font.size = Pt(24)
#Above is the cadet name
       tor = doc.add_paragraph()
       asa = tor.add_run('TO THE RANK OF CADET')
       asa.bold = True
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#Above is the "to the rank.." part
       perfect = doc.add_paragraph()
       sam = perfect.add_run(rank)
       sam.bold = True
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
       font = sam.font
       font.size = Pt(22)
#Above is the cadet rank part
       inair = doc.add_paragraph()
       amy = inair.add_run('IN THE AIR FORCE JUNIOR RESERVE OFFICERS TRAINING CORPS')
       amy.bold = True
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#above is the 'In the air...." part
       pif = doc.add_paragraph("Concurrent with this appointment is the charge to carefully and diligently discharge all duties of the rank to which appointed by doing and performing all manner of things thereunto pertaining. I do further charge and require all personnel of lesser grade to render respect and obedience to appropriate orders. The appointee is further charged to observe and follow such orders and directions as may be given by supervisors and instructors acting according to the rules governing the Air Force Junior Reserve Officers Training Corps.", style='Normal')
       paragraph_format = pif.paragraph_format
       paragraph_format.space_after = Pt(0)
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
#Text above is the big paragraph thingy
       nummy = name.rfind(' ')
       nummy = nummy
       name1 = name[nummy:]
       name2 = name[:nummy]
       name = name1 + ',' + name2
       name.replace(" ", '')
       doc.add_picture(cwd + '/CaptureROTCDoc - Copy.png')
       docname = name + rank + '.docx'
       doc.save(docname)
####################################
       docxtopdf(docname)#Delete/disable to remove PDF function
####################################
def pdfcompiler(pdfs):
       merger = PdfFileMerger()
       for pdf in pdfs:
              merger.append(pdf)
       name = input("Please give a name for the combined PDF file (text only, no '.pdf'): ", width = Inches(10))
       merger.write(name + ".pdf")
pdfs = []
print("Starting.....")
print("Certificate Promotion Maker Loaded_")
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet Airman. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = str(input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: "))
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Airman')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet Airman First Class. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Airman First Class')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet Senior Airman. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Senior Airman')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Staff Sergeant. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Staff Sergeant')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet Technical Sergeant. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Technical Sergeant')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet Master Sergeant. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Master Sergeant')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet Senior Master Sergeant. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Senior Master Sergeant')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet Chief Master Sergeant. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Cadet Chief Master Sergeant')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet Second Lieutenant. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Second Lieutenant')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet First Lieutenant. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'First Lieutenant')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet Captain. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Captain')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet Major. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Major')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet Lieutenant Colonel. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Lieutenant Colonel')
while 1 == 1:
       print("Now going over cadets being promoted to the rank of Cadet Colonel. At any time type 'done' into the prompt to move on.")
       while 1 == 1:
              name = input("Please enter the full name of the cadet as it should come up on the all of the certificates or 'done' to move on: ")
              print('{} is what you entered. Is this correct? (Y/N)'.format(name))
              comm = input()
              if comm == 'y' or comm == 'Y':
                     break
              else:
                     print("We'll reenter it then.")
       if name.lower() == 'done':
              break
       else:
              docmaker(name, 'Colonel')
print()
pdfcompiler(pdfs)
print()
print("That concludes Promotion Certificate Maker. Press Enter To exit the program. All files make will show up at the location of the program. Thank you for using Promotion Certificate Maker!")
input()