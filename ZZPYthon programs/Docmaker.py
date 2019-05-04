import os
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
cwd = os.getcwd()
print(cwd)
def docmaker(name, rank):
       doc = Document()
       section = doc.sections[-1]
       section.orientation = WD_ORIENT.LANDSCAPE
       sections = doc.sections
       margin = 0.5
       for section in sections:
              section.top_margin = Inches(margin)
              section.bottom_margin = Inches(margin)
              section.left_margin = Inches(margin)
              section.right_margin = Inches(margin)
       style = doc.styles['Normal']
       font = style.font
       font.name = 'Times New Roman'
       font.size = Pt(14)
       new_width, new_height = section.page_height, section.page_width
       section.orientation = WD_ORIENT.LANDSCAPE
       section.page_width = new_width
       section.page_height = new_height
#code above sets page as landscape and sets margins, plus style "normal"
       doc.add_picture(cwd + '/AfJROTC EMBLEM.gif', width=Inches(1.75))
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#Code above sets picture
       cert = doc.add_paragraph()
       tim = cert.add_run('Certificate of Promotion')
       tim.bold = True
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
       font = tim.font
       font.size = Pt(24)
       font.color.rgb = RGBColor(31, 73, 125)
#That adds the "Certificate of promotion" part
       rec = doc.add_paragraph()
       bob = rec.add_run('IN RECOGNITION OF DEMONSTRATED FIDELITY AND ABILITIES I HEREBY APPOINT')
       bob.bold = True
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#Above is the "In recognition..." part
       sorry = doc.add_paragraph()
       geo = sorry.add_run(name)
       geo.bold = True
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
       font = geo.font
       font.size = Pt(24)
#Above is the cadet name
       tor = doc.add_paragraph()
       asa = tor.add_run('TO THE RANK OF CADET')
       asa.bold = True
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#Above is the "to the rank.." part
       perfect = doc.add_paragraph()
       sam = perfect.add_run(rank)
       sam.bold = True
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
       font = sam.font
       font.size = Pt(22)
#Above is the cadet rank part
       inair = doc.add_paragraph()
       amy = inair.add_run('IN THE AIR FORCE JUNIOR RESERVE OFFICERS TRAINING CORPS')
       amy.bold = True
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#above is the 'In the air...." part
       pif = doc.add_paragraph("Concurrent with this appointment is the charge to carefully and diligently discharge all duties of the rank to which appointed by doing and performing all manner of things thereunto pertaining. I do further charge and require all personnel of lesser grade to render respect and obedience to appropriate orders. The appointee is further charged to observe and follow such orders and directions as may be given by supervisors and instructors acting according to the rules governing the Air Force Junior Reserve Officers Training Corps.", style='Normal')
       paragraph_format = pif.paragraph_format
       paragraph_format.space_after = Pt(0)
       last_paragraph = doc.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
#Text above is the big paragraph thingy
       nummy = name.rfind(' ')
       nummy = nummy
       name1 = name[nummy:]
       name2 = name[:nummy]
       name = name1 + ',' + name2
       name.replace(" ", '')
       doc.add_picture(cwd + '/CaptureROTCDoc - Copy.png')
       docname = name + rank + '.docx'
       doc.save(docname)
docmaker('Arick Hauschild', 'Colonel')